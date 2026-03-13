import re
from datetime import date

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from datetime import datetime
from pathlib import Path

class DocBuilder:
    MINISTRY = "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РФ"
    UNIVERSITY = (
        "филиал федерального государственного бюджетного образовательного "
        "учреждения высшего образования «НИУ МЭИ» в г. Смоленске"
    )
    DEPARTMENT = "Кафедра вычислительной техники"
    DEFAULT_CITY = "Смоленск"

    def __init__(
        self,
        subject: str,
        work_type: str = "ОТЧЁТ",
        work_title: str = "",
        student: str = "",
        group: str = "",
        teacher: str = "",
        variant: int | None = None,
        city: str | None = None,
        year: int | None = None,
    ):
        if not subject.strip():
            raise ValueError("Поле subject обязательно")

        self.doc = Document()
        self.state = {
            "heading": [0] * 5,
            "table": 0,
            "figure": 0,
        }

        self.meta = {
            "subject": subject.strip(),
            "work_type": work_type.strip() or "ОТЧЁТ",
            "work_title": work_title.strip(),
            "student": student.strip(),
            "group": group.strip(),
            "teacher": teacher.strip(),
            "variant": variant,
            "city": (city or self.DEFAULT_CITY).strip(),
            "year": year or date.today().year,
        }

        self._setup_page()
        self._setup_styles()
        self._add_title_page()

    def _setup_page(self):
        section = self.doc.sections[0]
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    def _setup_styles(self):
        normal = self.doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(14)

        pf = normal.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(1.25)

        for level in range(1, 6):
            style = self.doc.styles[f"Heading {level}"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(14)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)

            hpf = style.paragraph_format
            hpf.first_line_indent = Cm(0)
            hpf.space_before = Pt(12)
            hpf.space_after = Pt(6)
            hpf.keep_with_next = True
            hpf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def _center_text(self, text: str, bold: bool = False, size: int = 14):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)

        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold

    def _empty_lines(self, count: int):
        for _ in range(count):
            self.doc.add_paragraph("")

    def _add_title_page(self):
        m = self.meta

        self._center_text(self.MINISTRY, bold=True)
        self._center_text(self.UNIVERSITY)
        self.doc.add_paragraph("")
        self._center_text(self.DEPARTMENT)

        self._empty_lines(3)

        self._center_text(m["work_type"], bold=True, size=16)

        if m["work_title"]:
            self._center_text(m["work_title"], bold=True)

        self._center_text(f'по дисциплине «{m["subject"]}»')

        self._empty_lines(6)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)

        lines = []
        if m["group"]:
            lines.append(f"Группа: {m['group']}")
        if m["student"]:
            lines.append(f"Студент: {m['student']}")
        if m["teacher"]:
            lines.append(f"Преподаватель: {m['teacher']}")
        if m["variant"] is not None:
            lines.append(f"Вариант: №{m['variant']}")

        run = p.add_run("\n".join(lines))
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        self._empty_lines(6)
        self._center_text(f'{m["city"]}, {m["year"]} г.')
        self.doc.add_page_break()

    def add(self, block):
        block.render(self)
        return self

    def _sanitize(self, text: str) -> str:
        """Удаляет символы, которые нельзя использовать в имени файла."""
        return re.sub(r'[\\/*?:"<>|]', "", text)

    def _generate_filename(self):
        m = self.meta

        parts = []

        if m["work_type"]:
            parts.append(m["work_type"])

        parts.append(m["subject"])

        if m["student"]:
            parts.append(m["student"])

        if m["group"]:
            parts.append(m["group"])

        name = " ".join(parts)
        name = self._sanitize(name)

        return f"{name}.docx"

    def save(self, filename=None):
        build_dir = Path("build")
        build_dir.mkdir(exist_ok=True)

        if filename is None:
            filename = self._generate_filename()

        path = build_dir / filename

        self.doc.save(path)
        return str(path)