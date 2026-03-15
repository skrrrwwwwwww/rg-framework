from abc import ABC, abstractmethod

from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class Block(ABC):
    @abstractmethod
    def render(self, builder):
        pass


class Paragraph(Block):
    def __init__(self, text, align="justify", indent=True, bold=False, size=14):
        self.text = text
        self.align = align
        self.indent = indent
        self.bold = bold
        self.size = size

    def render(self, builder):
        p = builder.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25) if self.indent else Cm(0)
        p.paragraph_format.keep_together = True
        p.paragraph_format.widow_control = True

        if self.align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif self.align == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run(self.text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(self.size)
        run.bold = self.bold


class Heading(Block):
    def __init__(self, text, level=1, numbered=True):
        if not 1 <= level <= 5:
            raise ValueError("level должен быть от 1 до 5")
        self.text = text.rstrip(".")
        self.level = level
        self.numbered = numbered

    def render(self, builder):
        if self.numbered:
            idx = self.level - 1
            builder.state["heading"][idx] += 1

            for i in range(idx + 1, len(builder.state["heading"])):
                builder.state["heading"][i] = 0

            prefix = ".".join(str(x) for x in builder.state["heading"][:self.level])
            text = f"{prefix} {self.text}"
        else:
            text = self.text

        h = builder.doc.add_heading(text, level=self.level)
        for run in h.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.keep_together = True
        h.paragraph_format.widow_control = True


class TableOfContents(Block):
    def __init__(self, title="СОДЕРЖАНИЕ", levels="1-5", page_break_after=True):
        self.title = title
        self.levels = levels
        self.page_break_after = page_break_after

    def render(self, builder):
        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.keep_with_next = True

        run = p.add_run(self.title)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True

        builder.doc.add_paragraph("")

        p = builder.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)

        run = p.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f'TOC \\o "{self.levels}" \\h \\z \\u'

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")

        text = OxmlElement("w:t")
        text.text = "Оглавление будет создано после обновления поля в Word."

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(text)
        run._r.append(fld_end)

        if self.page_break_after:
            builder.doc.add_page_break()


class NumberedList(Block):
    def __init__(self, items):
        self.items = items

    def render(self, builder):
        for i, item in enumerate(self.items, start=1):
            p = builder.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(0.63)
            pf.first_line_indent = Cm(-0.63)
            pf.keep_together = True
            pf.widow_control = True

            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run = p.add_run(f"{i}. {item}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)


class Table(Block):
    def __init__(self, title, data):
        self.title = title
        self.data = data

    def render(self, builder):
        builder.state["table"] += 1
        number = builder.state["table"]

        cap = builder.doc.add_paragraph(f"Таблица {number} – {self.title}")
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.keep_with_next = True
        cap.paragraph_format.keep_together = True
        cap.paragraph_format.widow_control = True
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if not self.data:
            return

        rows = len(self.data)
        cols = max(len(row) for row in self.data)

        table = builder.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(self.data):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                cell.text = str(value)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)


class ImagePlaceholder(Block):
    def __init__(self, title, height_lines=6):
        self.title = title
        self.height_lines = height_lines

    def render(self, builder):
        builder.state["figure"] += 1
        number = builder.state["figure"]

        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True

        p.add_run("\n" * self.height_lines)
        p.add_run("[МЕСТО ДЛЯ ИЗОБРАЖЕНИЯ / СКРИНШОТА]")
        p.add_run("\n" * self.height_lines)

        cap = builder.doc.add_paragraph(f"Рисунок {number} – {self.title}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.keep_together = True
        cap.paragraph_format.widow_control = True


class Formula(Block):
    def __init__(self, formula_text, number=None, explanation=None):
        self.formula_text = formula_text
        self.number = number
        self.explanation = explanation

    def render(self, builder):
        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.keep_together = True
        p.paragraph_format.widow_control = True

        if self.number is None:
            p.add_run(self.formula_text)
        else:
            tabs = p.paragraph_format.tab_stops
            tabs.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            p.add_run(self.formula_text)
            p.add_run("\t")
            p.add_run(f"({self.number})")

        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

        if self.explanation:
            exp = builder.doc.add_paragraph(self.explanation)
            exp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            exp.paragraph_format.first_line_indent = Cm(1.25)
            exp.paragraph_format.keep_together = True
            exp.paragraph_format.widow_control = True


class PageBreak(Block):
    def render(self, builder):
        builder.doc.add_page_break()