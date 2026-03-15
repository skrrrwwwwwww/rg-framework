import re
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocBuilder:
    MINISTRY = "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РФ"
    UNIVERSITY = (
        "филиал федерального государственного бюджетного образовательного "
        "учреждения высшего образования «Национальный исследовательский "
        "университет «МЭИ» в г. Смоленске"
    )
    DEPARTMENT = "Кафедра вычислительной техники"
    DEFAULT_CITY = "Смоленск"

    def __init__(
        self,
        subject: str,
        work_type: str = "ОТЧЁТ",
        work_title: str = "",
        student: str = "Шикарев Иван А",
        group: str = "ПО1-23",
        teacher: str = "",
        variant: int | None = 17,
        city: str | None = None,
        year: int | None = None,
        subject_short: str | None = None,
    ):
        if not subject.strip():
            raise ValueError("Поле subject обязательно")

        self.doc = Document()
        self.state = {
            "heading": [0] * 5,
            "table": 0,
            "figure": 0,
            "formula": 0,
        }

        self.meta = {
            "subject": subject.strip(),
            "subject_short": (subject_short or subject).strip(),
            "work_type": work_type.strip() or "ОТЧЁТ",
            "work_title": work_title.strip(),
            "student": self._format_name_for_doc(student),
            "group": group.strip(),
            "teacher": teacher.strip(),
            "variant": variant,
            "city": (city or self.DEFAULT_CITY).strip(),
            "year": year or date.today().year,
        }

        # Сохраняем сырые данные для имени файла
        self.raw_data = {
            "student": student.strip(),
            "teacher": teacher.strip(),
        }

        # Форматируем для документа (Фамилия И.О.)
        student_for_doc = self._format_name_for_doc(student)

        self._setup_page()
        self._setup_styles()
        self._enable_hyphenation()
        self._add_title_page()

    def _setup_page(self):
        section = self.doc.sections[0]
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    def _setup_styles(self):
        normal = self.doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(14)

        pf = normal.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.first_line_indent = Cm(1.25)
        pf.widow_control = True
        pf.keep_together = True

        for level in range(1, 6):
            style = self.doc.styles[f"Heading {level}"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(14)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)

            hpf = style.paragraph_format
            hpf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            hpf.first_line_indent = Cm(1.25)
            hpf.space_before = Pt(12)
            hpf.space_after = Pt(6)
            hpf.keep_with_next = True
            hpf.keep_together = True
            hpf.widow_control = True
            hpf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _enable_hyphenation(self):
        settings = self.doc.settings.element
        auto_hyphenation = OxmlElement("w:autoHyphenation")
        auto_hyphenation.set(qn("w:val"), "1")
        settings.append(auto_hyphenation)

    def _center_text(
            self,
            text: str,
            bold: bool = False,
            size: int = 14,
            space_before: int = 0,
            space_after: int = 0,
            line_spacing: str = "single"
    ):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.keep_together = True
        pf.widow_control = True
        # Выбор интервала
        if line_spacing == "one_point_five":
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:  # single по умолчанию
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)

        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold

    def _empty_lines(self, count: int):
        for _ in range(count):
            self.doc.add_paragraph("")

    def add(self, block):
        block.render(self)
        return self

    def _add_title_page(self):
        m = self.meta

        self._center_text(self.MINISTRY, bold=False, space_after=0, line_spacing="one_point_five")
        self._center_text(self.UNIVERSITY, space_after=36, line_spacing="one_point_five")
        self._center_text(self.DEPARTMENT, space_after=72, line_spacing="one_point_five")

        self._center_text(m["work_type"], bold=False, size=14, space_after=0, line_spacing="one_point_five")

        if m["work_title"]:
            self._center_text(m["work_title"], bold=True, space_after=12, line_spacing="one_point_five")

        self._center_text(f'по дисциплине «{m["subject"]}»', space_after=180, line_spacing="one_point_five")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.keep_together = True
        pf.widow_control = True
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_after = Pt(70)  # ← УМЕНЬШИТЬ С 120 ДО 60

        lines = []
        if m["group"]:
            lines.append(f"Группа: {m['group']}")
        if m["student"]:
            lines.append(f"Студент: {m['student']}")
        if m["teacher"]:
            lines.append(f"Преподаватель: {m['teacher']}")
        if m["variant"] is not None:
            lines.append(f"Вариант: №{m['variant']}")

        run = p.add_run("\n".join(lines))
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        self._center_text(f'{m["city"]}, {m["year"]} г.', line_spacing="single")
        self.doc.add_page_break()

    def _clear_paragraph(self, paragraph):
        p = paragraph._element
        for child in list(p):
            p.remove(child)

    def _add_page_field(self, paragraph):
        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)

    def add_page_numbers_bottom(self, hide_on_first_page=True):
        section = self.doc.sections[0]

        if hide_on_first_page:
            section.different_first_page_header_footer = True

        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        self._clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_page_field(p)

        if hide_on_first_page:
            first_footer = section.first_page_footer
            fp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
            self._clear_paragraph(fp)

        return self

    def _sanitize_filename_part(self, value: str) -> str:
        value = re.sub(r'[\\/*?:"<>|]', "", value)
        value = re.sub(r"\s+", " ", value).strip()
        return value

    def _shorten_subject(self, text: str) -> str:
        """Методы анализа данных → МАД"""

        words = text.split()

        letters = []
        for w in words:
            if w and w[0].isalpha():
                letters.append(w[0].upper())

        return "".join(letters)

    def _shorten_work_type(self, text: str) -> str:
        """Преобразует 'Лабораторная работа 1' → 'ЛБ1'"""

        text = text.lower()

        if "лабораторная" in text:
            import re
            number = re.search(r"\d+", text)
            if number:
                return f"ЛБ{number.group()}"
            return "ЛБ"

        if "практическая" in text:
            import re
            number = re.search(r"\d+", text)
            if number:
                return f"ПР{number.group()}"
            return "ПР"

        if "расчетно-графическая" in text:
            import re
            number = re.search(r"\d+", text)
            if number:
                return f"РГР{number.group()}"
            return "РГР"

        if "курсовая" in text:
            return "КР"

        return text.upper()

    def _format_name_for_doc(self, full_name: str) -> str:
        """Преобразует 'Шикарев Иван А' в 'Шикарев И.А.'"""
        if not full_name:
            return ""

        parts = full_name.strip().split()
        if len(parts) == 3:  # Фамилия Имя Отчество
            return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
        elif len(parts) == 2:  # Фамилия Имя
            return f"{parts[0]} {parts[1][0]}."
        else:
            return full_name

    def _format_name_for_filename(self, full_name: str) -> str:
        """Преобразует 'Иванов Иван Иванович' в 'Иванов Иван'"""
        if not full_name:
            return ""

        parts = full_name.strip().split()
        if len(parts) >= 2:
            return f"{parts[0]} {parts[1]}"
        return parts[0] if parts else ""

    def _build_default_filename(self) -> str:
        """Формирует имя файла: РГР ТОП Шикарев Иван ПО1-23.docx"""

        m = self.meta

        # Сокращаем тип работы
        work_type_short = self._shorten_work_type(m["work_type"])

        # Сокращаем дисциплину (первые буквы)
        subject_short = self._shorten_subject(m["subject"])

        student_for_filename = self._format_name_for_filename(self.raw_data["student"])

        # Собираем имя файла
        parts = [
            work_type_short,
            subject_short,
            student_for_filename,
            m["group"],
        ]

        # Убираем пустые части
        parts = [part for part in parts if part]

        # Если ничего не получилось, используем базовое имя
        if not parts:
            return "document.docx"

        # Объединяем части пробелами
        filename = " ".join(parts)

        # Очищаем от недопустимых символов
        filename = self._sanitize_filename_part(filename)

        return f"{filename}.docx"


    def save(self, filename: str | None = None, folder: str = "build"):
        build_dir = Path(folder)
        build_dir.mkdir(parents=True, exist_ok=True)

        if filename is None:
            filename = self._build_default_filename()

        path = build_dir / filename

        # Если файл существует и заблокирован, создаем уникальное имя
        counter = 1
        while True:
            try:
                # Пробуем сохранить
                self.doc.save(str(path))
                return str(path)
            except PermissionError:
                # Если файл заблокирован, добавляем номер
                name_without_ext = path.stem
                ext = path.suffix
                new_filename = f"{name_without_ext}_{counter}{ext}"
                path = build_dir / new_filename
                counter += 1
            except Exception as e:
                # Другие ошибки
                raise e