from abc import ABC, abstractmethod

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class Block(ABC):
    @abstractmethod
    def render(self, builder):
        pass


class Paragraph(Block):
    def __init__(self, text, align="justify", indent=True, bold=False, size=14):
        self.text = text
        self.align = align
        self.indent = indent
        self.bold = bold
        self.size = size

    def render(self, builder):
        p = builder.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25) if self.indent else Cm(0)

        if self.align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif self.align == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run(self.text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(self.size)
        run.bold = self.bold


class Heading(Block):
    def __init__(self, text, level=1, numbered=True):
        if not 1 <= level <= 5:
            raise ValueError("level должен быть от 1 до 5")
        self.text = text
        self.level = level
        self.numbered = numbered

    def render(self, builder):
        if self.numbered:
            idx = self.level - 1
            builder.state["heading"][idx] += 1

            for i in range(idx + 1, len(builder.state["heading"])):
                builder.state["heading"][i] = 0

            prefix = ".".join(str(x) for x in builder.state["heading"][:self.level])
            text = f"{prefix} {self.text}"
        else:
            text = self.text

        builder.doc.add_heading(text, level=self.level)


class Table(Block):
    def __init__(self, title, data):
        self.title = title
        self.data = data

    def render(self, builder):
        builder.state["table"] += 1
        number = builder.state["table"]

        cap = builder.doc.add_paragraph(f"Таблица {number} – {self.title}")
        cap.paragraph_format.first_line_indent = Cm(0)

        if not self.data:
            return

        rows = len(self.data)
        cols = max(len(row) for row in self.data)

        table = builder.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(self.data):
            for j, value in enumerate(row):
                table.cell(i, j).text = str(value)


class PageBreak(Block):
    def render(self, builder):
        builder.doc.add_page_break()